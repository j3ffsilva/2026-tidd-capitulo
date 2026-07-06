import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "capitulo.md"
OUT = "capitulo.docx"

def set_lang(doc, lang="pt-BR"):
    doc.core_properties.language = lang
    styles_element = doc.styles.element
    for rpr_default in styles_element.findall(qn('w:docDefaults') + '/' + qn('w:rPrDefault') + '/' + qn('w:rPr')):
        pass
    # Set language on Normal style run properties
    normal = doc.styles['Normal']
    rpr = normal.element.get_or_add_rPr()
    lang_el = rpr.find(qn('w:lang'))
    if lang_el is None:
        lang_el = OxmlElement('w:lang')
        rpr.append(lang_el)
    lang_el.set(qn('w:val'), lang)
    lang_el.set(qn('w:eastAsia'), lang)
    lang_el.set(qn('w:bidi'), lang)

def add_runs(paragraph, text):
    """Parse **bold** and *italic* markdown spans into runs."""
    pos = 0
    pattern = re.compile(r'(\*\*.+?\*\*|\*.+?\*)')
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith('**'):
            paragraph.add_run(token[2:-2]).bold = True
        else:
            paragraph.add_run(token[1:-1]).italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

def style_body_paragraph(p):
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(12)
    pf.space_before = Pt(0)

def style_reference_paragraph(p):
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_after = Pt(10)
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)

def main():
    with open(SRC, encoding="utf-8") as f:
        raw = f.read()

    lines = raw.split("\n")

    doc = Document()

    # base document font
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    set_lang(doc, "pt-BR")

    in_references = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        i += 1

        if not line.strip():
            continue

        if line.startswith("# "):
            title_text = line[2:].strip()
            h = doc.add_heading(level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_runs(h, title_text)
            continue

        if line.startswith("## "):
            heading_text = line[3:].strip()
            in_references = heading_text.strip().lower() == "referências"
            h = doc.add_heading(level=2)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_runs(h, heading_text)
            continue

        # Resumo / Palavras-chave / body / reference paragraphs
        p = doc.add_paragraph()
        add_runs(p, line)

        if in_references:
            style_reference_paragraph(p)
        else:
            style_body_paragraph(p)

    doc.save(OUT)
    print(f"Salvo em {OUT}")

if __name__ == "__main__":
    main()
